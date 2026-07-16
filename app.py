import streamlit as st
import os
from io import BytesIO
import fitz  # PyMuPDF
from docx import Document
import hr_db
import importlib
importlib.reload(hr_db)

import streamlit as st

# Initialize database
hr_db.init_db()
st.set_page_config(
    page_title="HR Yönetim Asistanı",
    layout="wide",
    page_icon="🏢",
    initial_sidebar_state="collapsed",
)

import base64
import io
import json
import html

def save_current_state_to_db():
    if st.session_state.get('current_record_id'):
        state_to_save = {
            "step": st.session_state.step,
            "tutanak_text": st.session_state.get('tutanak_text', ''),
            "tutanak_summary": st.session_state.get('tutanak_summary', ''),
            "savunma_istem_metni": st.session_state.get('savunma_istem_metni', ''),
            "savunma_text": st.session_state.get('savunma_text', ''),
            "karar_sonucu": st.session_state.get('karar_sonucu', ''),
            "secilen_karar": st.session_state.get('secilen_karar', ''),
            "nihai_belge": st.session_state.get('nihai_belge', ''),
            "yonetmelik_text": st.session_state.get('yonetmelik_text', '')
        }
        hr_db.update_state(st.session_state.current_record_id, json.dumps(state_to_save))

import urllib.request

def get_base64_of_bin_file(bin_file):
    try:
        with open(bin_file, 'rb') as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except Exception:
        return ""

logo_path = "logo-display-transparent.png"
logo_b64 = get_base64_of_bin_file(logo_path)
logo_mark_path = "logo-mark-transparent.png"
logo_mark_b64 = get_base64_of_bin_file(logo_mark_path)

# --- GLOBAL BILGI CSS INJECTION ---
st.markdown(f"""
<style>
/* Ana Arka Plan ve Yazı Rengi */
.stApp {{
    background: #FAFAFA !important; /* Premium pearl white */
    color: #111827 !important; /* Sleek dark gray almost black */
}}

/* Sayfa üst boşluğunu azaltma (Daha yukarıdan başlama) */
.block-container {{
    padding-top: 2.5rem !important; /* Navbar'ın kesmemesi için 2.5rem ideal */
    padding-bottom: 2rem !important;
}}

/* Tüm fontlar daha şık */
html, body, .stApp {{
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
}}

/* Streamlit etiketlerinin koyu yapılması (Okunabilirlik Garantisi) */
.st-emotion-cache-10trncj, label, .stMarkdown p, .st-emotion-cache-1wivap2 {{
    color: #0033A0 !important; /* Logo Kurumsal Mavi */
    font-weight: 500 !important;
}}

/* Streamlit expander (Ağaç Yapısı) Başlıkları */
[data-testid="stExpander"] {{
    background-color: #FFFFFF !important;
    border: 1px solid #E5E7EB !important;
    border-radius: 12px !important;
    box-shadow: 0 1px 2px rgba(0,0,0,0.02) !important;
    margin-bottom: 8px !important;
}}

/* Expander Başlık Metinleri (Logo Mavisi) */
[data-testid="stExpander"] summary p {{
    color: #0033A0 !important;
    font-weight: 600 !important;
    font-size: 15px !important;
}}

[data-testid="stExpander"]:hover {{
    border-color: #F05A28 !important; /* Logo Turuncu Vurgu */
}}

/* Expander İç Gövdesi (İçerik Kısmı) - Marka Renkleriyle Bütünleşik */
[data-testid="stExpanderDetails"] {{
    background: linear-gradient(to bottom, #F4F8FC, #FFFFFF) !important; /* Çok hafif logo mavisi tonu */
    border-top: 2px solid #F05A28 !important; /* İnce turuncu marka çizgisi */
    border-radius: 0 0 12px 12px !important;
    padding-top: 20px !important;
}}

/* Flat Design Konteynerlar ve Kartlar */
.st-emotion-cache-1wmy9hl, .st-emotion-cache-16txtl3, .stFileUploader, .stTextArea {{
    background: #FFFFFF !important;
    border: 1px solid #F3F4F6 !important;
    border-radius: 16px !important; /* Modern kavisler */
    padding: 24px !important;
    box-shadow: 0 4px 20px rgba(0,0,0,0.02) !important;
    transition: all 0.3s ease !important;
}}

.st-emotion-cache-1wmy9hl:hover, .st-emotion-cache-16txtl3:hover, .stFileUploader:hover {{
    box-shadow: 0 8px 30px rgba(0,0,0,0.04) !important;
    transform: translateY(-2px) !important;
}}

/* Premium AI Buttons (Logo Colors) */
.stButton > button {{
    background: linear-gradient(135deg, #0033A0 0%, #0055FF 100%) !important; /* Logo Koyu Mavi'den Açık Mavi'ye */
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 10px !important; 
    font-weight: 600 !important;
    letter-spacing: 0.3px !important;
    padding: 0.7rem 1.5rem !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 4px 12px rgba(0, 51, 160, 0.2) !important;
}}

.stButton > button:hover {{
    background: linear-gradient(135deg, #F05A28 0%, #F57245 100%) !important; /* Hover'da Logo Turuncusu */
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 16px rgba(240, 90, 40, 0.3) !important;
    color: #FFFFFF !important;
    border: none !important;
}}

/* Başlık ve Metin Rengi Düzeltmeleri */
h1, h2, h3, h4, h5, h6 {{
    color: #0033A0 !important; /* Logo Kurumsal Mavi */
    font-weight: 700 !important;
    letter-spacing: -0.5px !important;
}}

.main-app-title {{
    color: #F05A28 !important; /* Ana Başlık Turuncu Override */
}}

/* Input Alanları */
.stTextArea textarea, .stTextInput input, .stSelectbox > div > div {{
    background: #F9FAFB !important;
    border: 1px solid #E5E7EB !important;
    color: #1F2937 !important;
    border-radius: 10px !important;
    transition: all 0.2s ease !important;
}}

.stTextArea textarea:focus, .stTextInput input:focus, .stSelectbox > div > div:focus {{
    border-color: #F05A28 !important; /* Odaklanmada Turuncu */
    box-shadow: 0 0 0 3px rgba(240, 90, 40, 0.15) !important;
    background: #FFFFFF !important;
}}

/* Mobil bankacılık esintili, özgün uygulama katmanı */
:root {{
    --brand-navy: #0033A0;
    --brand-blue: #0B5FFF;
    --brand-orange: #F05A28;
    --ink: #10213E;
    --muted: #64748B;
    --line: #E3EAF5;
}}

.stApp {{
    background: linear-gradient(180deg, #EEF4FF 0, #F7F9FD 330px, #F4F7FC 100%) !important;
    color: var(--ink) !important;
}}

.block-container {{
    max-width: 1220px !important;
    padding-top: 1.5rem !important;
    padding-bottom: 3rem !important;
}}

label, [data-testid="stWidgetLabel"] p {{
    color: var(--ink) !important;
    font-weight: 650 !important;
}}

[data-testid="stExpander"] {{
    background: #FFFFFF !important;
    border: 1px solid var(--line) !important;
    border-radius: 18px !important;
    box-shadow: 0 8px 24px rgba(15, 45, 102, 0.05) !important;
    margin-bottom: 12px !important;
    overflow: hidden !important;
}}

[data-testid="stExpander"] summary p {{
    color: var(--ink) !important;
    font-weight: 750 !important;
    font-size: 0.96rem !important;
}}

[data-testid="stExpander"]:hover {{
    border-color: #AFC5F1 !important;
    box-shadow: 0 12px 28px rgba(15, 45, 102, 0.09) !important;
}}

[data-testid="stExpanderDetails"] {{
    background: linear-gradient(180deg, #F8FBFF 0%, #FFFFFF 140px) !important;
    border-top: 1px solid var(--line) !important;
    padding-top: 18px !important;
}}

.stFileUploader {{
    background: #F8FAFE !important;
    border: 1px dashed #B6C8E8 !important;
    border-radius: 14px !important;
    padding: 16px !important;
}}

.stTextArea textarea, .stTextInput input, .stSelectbox > div > div {{
    background: #FFFFFF !important;
    border: 1px solid #D5E0F0 !important;
    color: var(--ink) !important;
    border-radius: 12px !important;
    min-height: 46px !important;
}}

.stTextArea textarea:focus, .stTextInput input:focus, .stSelectbox > div > div:focus-within {{
    border-color: var(--brand-blue) !important;
    box-shadow: 0 0 0 3px rgba(11, 95, 255, 0.12) !important;
}}

.stButton > button {{
    background: linear-gradient(135deg, var(--brand-navy) 0%, var(--brand-blue) 100%) !important;
    border-radius: 12px !important;
    font-weight: 700 !important;
    min-height: 44px !important;
    padding: 0.62rem 1rem !important;
    box-shadow: 0 6px 14px rgba(0, 51, 160, 0.18) !important;
}}

.stButton > button:hover {{
    background: linear-gradient(135deg, #0A2A78 0%, #0348C9 100%) !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 9px 18px rgba(0, 51, 160, 0.24) !important;
}}

h1, h2, h3, h4, h5, h6 {{ color: var(--ink) !important; }}

.app-shell {{
    background: linear-gradient(135deg, #002A83 0%, #084ED6 100%);
    border-radius: 22px;
    padding: 22px 26px;
    color: #FFFFFF;
    box-shadow: 0 14px 32px rgba(0, 51, 160, 0.2);
    margin-bottom: 18px;
}}

.app-logo-card {{
    background: #FFFFFF;
    border: 1px solid rgba(255,255,255,0.55);
    border-radius: 16px;
    min-width: 148px;
    padding: 10px 14px;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    box-shadow: 0 8px 18px rgba(0, 28, 91, 0.2);
}}

.app-logo-card img {{
    display: block;
    width: 124px;
    max-height: 48px;
    object-fit: contain;
}}

.app-eyebrow {{ font-size: 0.72rem; font-weight: 750; letter-spacing: 0.12em; opacity: 0.78; text-transform: uppercase; }}
.app-welcome {{ font-size: 1.55rem; font-weight: 800; margin: 4px 0; letter-spacing: -0.6px; }}
.app-subtitle {{ font-size: 0.88rem; opacity: 0.82; margin: 0; }}

.summary-card {{
    background: #FFFFFF;
    border: 1px solid #E2EAF6;
    border-radius: 18px;
    padding: 18px;
    min-height: 114px;
    box-shadow: 0 8px 22px rgba(15, 45, 102, 0.05);
}}
.summary-label {{ color: #64748B; font-size: 0.72rem; font-weight: 750; letter-spacing: 0.05em; text-transform: uppercase; }}
.summary-value {{ color: #10213E; font-size: 1.8rem; font-weight: 800; line-height: 1.2; margin-top: 8px; }}
.summary-detail {{ color: #64748B; font-size: 0.78rem; margin-top: 7px; }}
.section-intro {{ display: flex; align-items: center; justify-content: space-between; margin: 24px 2px 12px; }}
.section-intro h3 {{ margin: 0; font-size: 1.08rem; }}
.section-intro span {{ color: #64748B; font-size: 0.82rem; }}
.logo-caption {{ color: #0033A0; font-size: 0.64rem; font-weight: 750; letter-spacing: 0.06em; margin-top: 4px; text-align: center; text-transform: uppercase; }}

[data-testid="stSidebar"] {{ background: #F7F9FD !important; border-right: 1px solid var(--line); }}
[data-testid="stSidebar"] [data-testid="stExpander"] {{ border-radius: 14px !important; box-shadow: none !important; }}

.login-hero {{ margin: 3rem auto 1.25rem; max-width: 560px; text-align: center; }}
.login-brand {{ background: #FFFFFF; border: 1px solid #E2EAF6; border-radius: 24px; box-shadow: 0 16px 38px rgba(15, 45, 102, 0.08); padding: 26px 24px 10px; }}

@media (max-width: 700px) {{
    .block-container {{ padding: 0.9rem 0.9rem 2rem !important; }}
    .app-shell {{ border-radius: 18px; padding: 18px; }}
    .app-logo-card {{ min-width: 105px; padding: 8px 10px; border-radius: 13px; }}
    .app-logo-card img {{ width: 92px; max-height: 38px; }}
    .logo-caption {{ font-size: 0.55rem; }}
    .app-welcome {{ font-size: 1.25rem; }}
    .summary-card {{ min-height: 96px; padding: 14px; border-radius: 15px; }}
    .summary-value {{ font-size: 1.45rem; }}
    [data-testid="stExpander"] {{ border-radius: 15px !important; }}
    .login-hero {{ margin-top: 1.3rem; }}
}}

/* Sade görünüm: içerik ve mevcut süreç akışı ön planda */
.stApp {{ background: #F7F8FA !important; color: #1F2937 !important; }}
.block-container {{ max-width: 1120px !important; padding-top: 0.5rem !important; padding-bottom: 2.5rem !important; }}
html, body, .stApp {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif !important; }}
.stMarkdown p {{ color: #334155 !important; font-weight: 400 !important; }}
label, [data-testid="stWidgetLabel"] p {{ color: #334155 !important; font-weight: 600 !important; }}
h1, h2, h3, h4, h5, h6 {{ color: #0033A0 !important; letter-spacing: normal !important; }}

.app-shell {{
    background: #FFFFFF;
    border: 1px solid #E5E7EB;
    border-radius: 12px;
    padding: 18px 20px;
    color: #1F2937;
    box-shadow: 0 2px 8px rgba(15, 23, 42, 0.04);
    margin-bottom: 16px;
}}
.app-eyebrow {{ color: #64748B; font-size: 0.72rem; font-weight: 700; letter-spacing: 0.04em; opacity: 1; text-transform: uppercase; }}
.app-welcome {{ color: #0033A0; font-size: 1.35rem; font-weight: 750; margin: 4px 0; letter-spacing: normal; }}
.app-shell .app-subtitle {{ color: #64748B !important; font-size: 0.86rem; opacity: 1; margin: 0; }}
.app-logo-card {{ background: transparent !important; border: none !important; border-radius: 0 !important; min-width: 0; padding: 0 !important; box-shadow: none !important; }}
.app-logo-card img {{ width: 112px; max-height: 44px; }}
.logo-caption {{ color: #64748B; font-size: 0.58rem; font-weight: 600; letter-spacing: 0.02em; margin-top: 3px; }}

.summary-card {{ background: #FFFFFF; border: 1px solid #E5E7EB; border-radius: 12px; padding: 14px; min-height: 98px; box-shadow: none; }}
.summary-label {{ color: #64748B; font-size: 0.7rem; font-weight: 700; letter-spacing: normal; text-transform: uppercase; }}
.summary-value {{ color: #1F2937; font-size: 1.55rem; font-weight: 750; line-height: 1.2; margin-top: 6px; }}
.summary-detail {{ color: #64748B; font-size: 0.74rem; margin-top: 5px; }}
.section-intro {{ margin: 20px 2px 10px; }}
.section-intro h3 {{ color: #0033A0 !important; font-size: 1rem; }}

[data-testid="stExpander"] {{ background: #FFFFFF !important; border: 1px solid #E5E7EB !important; border-radius: 12px !important; box-shadow: none !important; margin-bottom: 8px !important; }}
[data-testid="stExpander"] summary p {{ color: #0033A0 !important; font-size: 0.94rem !important; font-weight: 650 !important; }}
[data-testid="stExpander"]:hover {{ border-color: #B8C7E3 !important; box-shadow: none !important; }}
[data-testid="stExpanderDetails"] {{ background: #FFFFFF !important; border-top: 1px solid #E5E7EB !important; padding-top: 16px !important; }}

.stButton > button {{ background: #0033A0 !important; border: 1px solid #0033A0 !important; border-radius: 8px !important; box-shadow: none !important; font-weight: 650 !important; min-height: 42px !important; padding: 0.55rem 1rem !important; transform: none !important; }}
.stButton > button:hover {{ background: #002878 !important; border-color: #002878 !important; box-shadow: none !important; transform: none !important; }}
.stTextArea textarea, .stTextInput input, .stSelectbox > div > div {{ background: #FFFFFF !important; border: 1px solid #D1D5DB !important; border-radius: 8px !important; min-height: 42px !important; }}
.stTextArea textarea:focus, .stTextInput input:focus, .stSelectbox > div > div:focus-within {{ border-color: #0033A0 !important; box-shadow: 0 0 0 2px rgba(0, 51, 160, 0.12) !important; }}
.stFileUploader {{ background: #FFFFFF !important; border: 1px dashed #CBD5E1 !important; border-radius: 10px !important; padding: 12px !important; }}

@media (max-width: 700px) {{
    .block-container {{ padding: 0.35rem 0.75rem 2rem !important; }}
    .app-shell {{ border-radius: 10px; padding: 14px; }}
    .app-logo-card {{ min-width: 0; padding: 0 !important; }}
    .app-logo-card img {{ width: 82px; max-height: 34px; }}
    .app-welcome {{ font-size: 1.15rem; }}
    .summary-card {{ min-height: 88px; padding: 12px; border-radius: 10px; }}
    .summary-value {{ font-size: 1.3rem; }}
}}

/* Giriş ekranı: üst alana yakın, sade ve logo odaklı */
.login-hero {{ margin: 0 auto 0.6rem; max-width: 620px; padding-top: 64px; }}
.login-brand {{ background: transparent !important; border: none !important; border-radius: 0 !important; padding: 0 24px 4px; box-shadow: none !important; }}
.login-brand {{ overflow: visible !important; }}
.login-brand img {{ display: block; width: 270px !important; max-width: 100% !important; height: auto !important; object-fit: contain !important; margin: 0 auto 2px !important; }}
.login-brand h2 {{ color: #0033A0 !important; font-family: Arial, Helvetica, sans-serif !important; font-size: 1.36rem !important; font-weight: 800 !important; letter-spacing: -0.03em !important; line-height: 1.05 !important; margin: 0 !important; }}
.login-brand p {{ color: #64748B !important; margin: 0 0 5px !important; }}
.login-title-accent {{ color: #F05A28 !important; }}
.login-hero {{ position: relative; }}
.login-watermark {{
    position: relative;
    width: min(25vw, 180px);
    margin: 10px auto 0;
    opacity: 0.10;
    pointer-events: none;
}}
.login-watermark img {{ display: block; width: 100%; height: auto; object-fit: contain; }}
.login-brand {{ background: transparent !important; }}
@media (max-width: 700px) {{
    .login-hero {{ margin-top: 0; padding-top: 44px; }}
    .login-brand {{ padding: 6px 14px 5px; }}
    .login-brand img {{ width: 250px !important; }}
    .login-watermark {{ width: 155px; margin-top: 8px; }}
}}
</style>
""", unsafe_allow_html=True)



# Session state initialization
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "role" not in st.session_state:
    st.session_state.role = ""
if "step" not in st.session_state:
    st.session_state.step = 0
if "tutanak_text" not in st.session_state:
    st.session_state.tutanak_text = ""
if "tutanak_summary" not in st.session_state:
    st.session_state.tutanak_summary = ""
if "savunma_istem_metni" not in st.session_state:
    st.session_state.savunma_istem_metni = ""
if "savunma_text" not in st.session_state:
    st.session_state.savunma_text = ""
if "karar_sonucu" not in st.session_state:
    st.session_state.karar_sonucu = ""
if "nihai_belge" not in st.session_state:
    st.session_state.nihai_belge = ""
if "yonetmelik_text" not in st.session_state:
    st.session_state.yonetmelik_text = ""


def login_screen():
    st.markdown(f"""
    <div class="login-hero">
      <div class="login-brand">
        <img src="data:image/png;base64,{logo_b64}" alt="ay.cx logosu">
        <h2 style="text-align:center; color:#10213E;"><span class="login-title-accent">HR</span> Yönetim Asistanı</h2>
        <div class="login-watermark" aria-hidden="true">
          <img src="data:image/png;base64,{logo_mark_b64}" alt="">
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1.35, 1])
    with col2:
        tab1, tab2 = st.tabs(["🔒 Giriş Yap", "📝 Yetki Talep Et"])
        
        with tab1:
            with st.form("login_form"):
                username = st.text_input("Kullanıcı Adı")
                password = st.text_input("Şifre", type="password")
                submit = st.form_submit_button("Giriş Yap")
                
                if submit:
                    if username and password:
                        user_data = hr_db.authenticate_user(username, password)
                        if user_data:
                            if user_data["status"] == "approved":
                                st.session_state.logged_in = True
                                st.session_state.username = username
                                st.session_state.role = user_data["role"]
                                hr_db.log_action(username, "Sisteme giriş yaptı.")
                                st.rerun()
                            else:
                                st.warning("Yetki talebiniz henüz onaylanmamış. Lütfen yöneticinin onaylamasını bekleyin.")
                        else:
                            st.error("Kullanıcı adı veya şifre hatalı!")
                    else:
                        st.warning("Lütfen kullanıcı adı ve şifre girin.")

        with tab2:
            with st.form("register_form"):
                st.info("Sisteme erişmek için yetki talebinizi oluşturun.")
                reg_username = st.text_input("Kullanıcı Adı Belirleyin")
                reg_password = st.text_input("Şifre Belirleyin", type="password")
                reg_submit = st.form_submit_button("Yetki Talep Et")
                
                if reg_submit:
                    if reg_username and reg_password:
                        success = hr_db.create_user(reg_username, reg_password)
                        if success:
                            st.success("Talebiniz iletildi! Yönetici onayından sonra giriş yapabilirsiniz.")
                        else:
                            st.error("Bu kullanıcı adı alınmış. Lütfen başka bir isim seçin.")
                    else:
                        st.warning("Lütfen kullanıcı adı ve şifre belirleyin.")

if not st.session_state.logged_in:
    login_screen()
    st.stop()

def extract_text_from_pdf(uploaded_file):
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

import glob

import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def set_text_preserve_style(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for i in range(1, len(paragraph.runs)):
        paragraph.runs[i].text = ""

def set_cell_preserve_style(cell, new_text):
    if cell.paragraphs:
        set_text_preserve_style(cell.paragraphs[0], new_text)
        # Diğer paragrafları temizle
        for i in range(1, len(cell.paragraphs)):
            cell.paragraphs[i].clear()
    else:
        cell.text = new_text

def create_docx(text, title="Belge"):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 1. Öncelik: app.py ile aynı klasördeki template.docx
    template_files = glob.glob(os.path.join(base_dir, "template.docx"))
    
    # 2. Öncelik: Eski sistem (HR klasörü altındaki Örnek.docx)
    if not template_files:
        hr_dir = os.path.join(base_dir, "..", "HR")
        template_files = glob.glob(os.path.join(hr_dir, "*rnek.docx"))
    
    doc = None
    if template_files:
        try:
            doc = Document(template_files[0])
            
            # P0: Tarih
            if len(doc.paragraphs) > 0:
                set_text_preserve_style(doc.paragraphs[0], f"Tarih: {datetime.datetime.now().strftime('%d.%m.%Y')}")
            
            # P1: Başlık
            if len(doc.paragraphs) > 1:
                spaced_title = " ".join(list(title.upper()))
                set_text_preserve_style(doc.paragraphs[1], spaced_title)
                
            # Table 0: Bilgi Tablosu
            if len(doc.tables) > 0:
                table = doc.tables[0]
                if len(table.rows) > 2:
                    set_cell_preserve_style(table.rows[0].cells[0], "Konu")
                    set_cell_preserve_style(table.rows[0].cells[1], f": {title}")
                    set_cell_preserve_style(table.rows[1].cells[0], "Tarih")
                    set_cell_preserve_style(table.rows[1].cells[1], f": {datetime.datetime.now().strftime('%d.%m.%Y')}")
                    set_cell_preserve_style(table.rows[2].cells[0], "İlgili Personel")
                    set_cell_preserve_style(table.rows[2].cells[1], ": Belge Tebellüğ Eden")
                    
            # P2 ve P3: Bölüm Başlıkları (Renkleri koruyarak güncelle)
            if len(doc.paragraphs) > 3:
                set_text_preserve_style(doc.paragraphs[2], "1. Belge / Personel Bilgileri")
                set_text_preserve_style(doc.paragraphs[3], "2. İstem ve Karar Detayları")
                
            # P4 ve P5: Eski test metinlerini sil (Gövdeyi temizle)
            if len(doc.paragraphs) > 5:
                p4 = doc.paragraphs[4]
                p4._element.getparent().remove(p4._element)
                p5 = doc.paragraphs[4] # P4 silinince P5 kayarak 4. index'e geldi
                p5._element.getparent().remove(p5._element)
                
                insertion_point = doc.paragraphs[4] # Yeni P6 (Bilginize sunulur.)
            else:
                insertion_point = None
                
            # Yapay Zeka Metnini İki Yana Yaslı (Justify) Olarak Enjekte Et
            for p_text in text.split('\n'):
                if p_text.strip():
                    if insertion_point:
                        new_p = insertion_point.insert_paragraph_before(p_text.strip())
                        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        new_p.paragraph_format.space_after = Pt(12)
                    else:
                        new_p = doc.add_paragraph(p_text.strip())
                        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        new_p.paragraph_format.space_after = Pt(12)
            
            # Alt kısımdaki 'test yazısı' ibarelerini temizle ve resmiyete dök
            for p in doc.paragraphs:
                if "test belgesinin" in p.text or "test belgesi" in p.text:
                    new_text = p.text.replace("test belgesinin", "belgenin").replace("test belgesi", "belge")
                    new_text = new_text.replace("içeriğini okudum ve teyit ettim", "içeriğini okudum ve bir nüshasını elden teslim aldım")
                    set_text_preserve_style(p, new_text)
                        
            io_stream = BytesIO()
            doc.save(io_stream)
            io_stream.seek(0)
            return io_stream
            
        except Exception as e:
            import streamlit as st
            import traceback
            st.error(f"Şablon enjeksiyon hatası: {e}\n\n{traceback.format_exc()}")
            print(f"Şablon enjeksiyon hatası: {e}")
            doc = None
            
    if doc is None:
        doc = Document()
        doc.add_heading(title, 0)
        for p in text.split('\n'):
            if p.strip():
                doc.add_paragraph(p)
    
    io_stream = BytesIO()
    doc.save(io_stream)
    io_stream.seek(0)
    return io_stream

display_username = html.escape(st.session_state.username)
st.markdown(f"""
<div class="app-shell">
    <div style="display:flex; align-items:center; justify-content:space-between; gap:18px;">
        <div>
            <div class="app-eyebrow">HR Yönetim Asistanı</div>
            <div class="app-welcome">Merhaba, {display_username}</div>
            <p class="app-subtitle">Disiplin süreçlerinize buradan güvenle devam edin.</p>
        </div>
        <div class="app-logo-card">
            <img src="data:image/png;base64,{logo_b64}" alt="ay.cx logosu">
            <div class="logo-caption">İK çalışma alanı</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# --- DASHBOARD SUMMARY (MOVED TO TOP) ---
stats = hr_db.get_summary_stats()
devam_eden_total = stats['tutanak'] + stats['bekleyen'] + stats['alindi']

summary_col1, summary_col2, summary_col3 = st.columns([1, 1.45, 1])
with summary_col1:
    st.markdown(f"""
    <div class="summary-card">
        <div class="summary-label">Toplam Dosya</div>
        <div class="summary-value">{stats['total']}</div>
        <div class="summary-detail">Sistemdeki tüm kayıtlar</div>
    </div>
    """, unsafe_allow_html=True)
with summary_col2:
    st.markdown(f"""
    <div class="summary-card">
        <div class="summary-label">Devam Eden Süreçler</div>
        <div class="summary-value" style="color:#D97706;">{devam_eden_total}</div>
        <div class="summary-detail">📝 {stats['tutanak']} tutanak &nbsp;•&nbsp; ⏳ {stats['bekleyen']} bekleyen &nbsp;•&nbsp; 🛡️ {stats['alindi']} alınan</div>
    </div>
    """, unsafe_allow_html=True)
with summary_col3:
    st.markdown(f"""
    <div class="summary-card">
        <div class="summary-label">Sonuçlanan</div>
        <div class="summary-value" style="color:#047857;">{stats['sonuc']}</div>
        <div class="summary-detail">Kararı tamamlanan dosyalar</div>
    </div>
    """, unsafe_allow_html=True)

# Çıkış Yap Butonu
col_logout1, col_logout2 = st.columns([6, 1])
with col_logout2:
    if st.button("Oturumu Kapat", key="logout_btn", use_container_width=True):
        hr_db.log_action(st.session_state.username, "Sistemden çıkış yaptı.")
        st.session_state.logged_in = False
        st.session_state.username = ""
        st.session_state.role = ""
        st.rerun()

# --- ADMIN PANEL ---
if st.session_state.role == "admin":
    with st.sidebar:
        st.markdown("<h3 style='text-align: center; color: #0033A0; margin-bottom: 20px;'>🛠️ Yönetici Menüsü</h3>", unsafe_allow_html=True)
        with st.expander("🔒 Yetki Yönetimi", expanded=False):
            st.write("Sisteme girmek için yetki talep eden kullanıcıları buradan onaylayabilir veya silebilirsiniz.")
            pending_users = hr_db.get_pending_users()
            if not pending_users:
                st.info("Onay bekleyen yeni yetki talebi bulunmuyor.")
            else:
                for u in pending_users:
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write(f"👤 **{u['username']}**")
                        st.write(f"({u['created_at'][:16]})")
                    with col2:
                        if st.button("✅ Onay", key=f"approve_{u['id']}"):
                            hr_db.approve_user(u['id'])
                            hr_db.log_action(st.session_state.username, f"Kullanıcı yetkisini onayladı: {u['username']}")
                            st.success(f"'{u['username']}' başarıyla onaylandı!")
                            st.rerun()
                        if st.button("❌ Red", key=f"reject_{u['id']}"):
                            hr_db.reject_user(u['id'])
                            hr_db.log_action(st.session_state.username, f"Kullanıcı talebini reddetti: {u['username']}")
                            st.warning(f"'{u['username']}' talebi silindi.")
                            st.rerun()
            
            st.markdown("---")
            st.subheader("Mevcut Kullanıcılar")
            approved_users = hr_db.get_approved_users()
            if not approved_users:
                st.info("Sistemde admin dışında aktif kullanıcı yok.")
            else:
                for u in approved_users:
                    st.write(f"👤 **{u['username']}** (Kayıt: {u['created_at'][:10]})")
                    if st.button("⛔ Yetkiyi Kaldır", key=f"revoke_{u['id']}", use_container_width=True):
                        hr_db.reject_user(u['id'])
                        hr_db.log_action(st.session_state.username, f"Kullanıcı yetkisini kaldırdı: {u['username']}")
                        st.error(f"'{u['username']}' yetkisi kaldırıldı!")
                        st.rerun()
                        
        with st.expander("📜 Sistem Log Kayıtları", expanded=False):
            st.write("Son 100 işlem kaydı.")
            if st.button("Yenile", use_container_width=True):
                st.rerun()
            logs = hr_db.get_logs()
            if logs:
                import pandas as pd
                df = pd.DataFrame(logs)
                df['created_at'] = pd.to_datetime(df['created_at']).dt.strftime("%m-%d %H:%M")
                df = df[['created_at', 'username', 'action']]
                df.columns = ["Tarih", "Kullanıcı", "İşlem"]
                st.dataframe(df, use_container_width=True, hide_index=True)
            else:
                st.info("Henüz log yok.")
                
        with st.expander("🔑 Şifre Değiştir", expanded=False):
            with st.form("pwd_form"):
                new_pwd = st.text_input("Yeni Şifre", type="password")
                new_pwd2 = st.text_input("Yeni Şifre (Tekrar)", type="password")
                if st.form_submit_button("Güncelle", use_container_width=True):
                    if new_pwd and new_pwd == new_pwd2:
                        hr_db.change_password(st.session_state.username, new_pwd)
                        hr_db.log_action(st.session_state.username, "Kendi şifresini değiştirdi.")
                        st.success("Şifreniz güncellendi! Lütfen giriş yapın.")
                        st.session_state.logged_in = False
                        st.rerun()
                    elif new_pwd != new_pwd2:
                        st.error("Şifreler eşleşmiyor!")
                    else:
                        st.error("Şifre boş olamaz!")
                        
        

# Süreç kartları için görünüm kontrolü
st.markdown("""
<div class="section-intro">
  <div><h3>Disiplin Süreci</h3><span>İlgili kartı açarak işleminize devam edin.</span></div>
  <span>7 adım</span>
</div>
""", unsafe_allow_html=True)

col_exp1, col_exp2 = st.columns(2)
with col_exp1:
    if st.button("Tüm Adımları Göster", use_container_width=True):
        st.session_state.expand_all = True
        st.rerun()
with col_exp2:
    if st.button("Sade Görünüme Dön", use_container_width=True):
        st.session_state.expand_all = False
        st.session_state.step = 0
        st.rerun()

# ----------------- ADIM 1: YÖNETMELİK -----------------
with st.expander("01  •  Ön Hazırlık — Yönetmelik ve Kanunlar", expanded=(st.session_state.get('expand_all', False) or st.session_state.step == 1)):
    
    kb_dir = os.path.join(os.path.dirname(__file__), "knowledge_base")
    os.makedirs(kb_dir, exist_ok=True)
    
    permanent_files = glob.glob(os.path.join(kb_dir, "*.txt"))
    permanent_text = ""
    saved_file_names = []
    
    for pf in permanent_files:
        with open(pf, "r", encoding="utf-8") as f:
            permanent_text += f.read() + "\n\n"
        saved_file_names.append(os.path.basename(pf).replace(".txt", ""))
        
    st.write("Bu aşamada sisteme yüklediğiniz ana kanunlar ve yönetmelikler **kalıcı olarak** hafızaya kazınır. Uygulamayı kapatsanız bile silinmez. Her yeni olayda tekrar yüklemenize gerek kalmaz.")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📚 Kalıcı Hafıza (Demirbaşlar)")
        if saved_file_names:
            st.info("Kayıtlı Belgeler:\n- " + "\n- ".join(saved_file_names))
            if st.button("Kalıcı Hafızayı Temizle"):
                for pf in permanent_files:
                    os.remove(pf)
                st.session_state.yonetmelik_text = ""
                st.rerun()
        else:
            st.warning("Hafızada henüz kayıtlı kalıcı belge yok.")
            
        yonetmelik_files = st.file_uploader("Kalıcı Belge Ekle (İş Kanunu vb.)", type=["docx", "pdf"], accept_multiple_files=True, key="permanent")
        if yonetmelik_files and st.button("Kalıcı Olarak Kaydet"):
            with st.spinner("Belgeler kalıcı hafızaya kaydediliyor..."):
                for y_file in yonetmelik_files:
                    text = ""
                    if y_file.type == "application/pdf":
                        try:
                            pdf_bytes = y_file.getvalue()
                            pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                            extracted = ""
                            for page in pdf_doc:
                                extracted += page.get_text()
                            
                            if len(extracted.strip()) > 50:
                                text = f"--- Belge: {y_file.name} ---\n{extracted}"
                            else:
                                st.warning(f"{y_file.name} içindeki metin okunamadı. Bu sürüm OCR kullanmaz; metni manuel olarak ekleyin.")
                        except Exception as e:
                            st.error(f"{y_file.name} Okuma hatası: {e}")
                    else:
                        doc = Document(y_file)
                        text = f"--- Belge: {y_file.name} ---\n" + "\n".join([p.text for p in doc.paragraphs])
                    
                    if text:
                        save_path = os.path.join(kb_dir, f"{y_file.name}.txt")
                        with open(save_path, "w", encoding="utf-8") as f:
                            f.write(text)
            hr_db.log_action(st.session_state.username, f"{len(yonetmelik_files)} adet kalıcı mevzuat/kural dosyası ekledi.")
            st.success(f"{len(yonetmelik_files)} belge KALICI olarak kaydedildi!")
            st.rerun()

    with col2:
        st.subheader("📎 Olaya Özel İlave Belgeler")
        st.write("Eğer sadece bu tutanak için eklemek istediğiniz geçici bir belge varsa yükleyin. (Sistem kapandığında silinir)")
        gecici_files = st.file_uploader("Geçici Belge Seç", type=["docx", "pdf"], accept_multiple_files=True, key="temporary")
        if gecici_files and st.button("Geçici Olarak İşle"):
            with st.spinner("Geçici belgeler okunuyor..."):
                all_gecici = []
                for g_file in gecici_files:
                    text = ""
                    if g_file.type == "application/pdf":
                        pdf_doc = fitz.open(stream=g_file.getvalue(), filetype="pdf")
                        for page in pdf_doc:
                            text += page.get_text()
                        if len(text.strip()) < 50:
                            st.warning(f"{g_file.name} içindeki metin okunamadı. Bu sürüm OCR kullanmaz; metni manuel olarak ekleyin.")
                    else:
                        doc = Document(g_file)
                        text = "\n".join([p.text for p in doc.paragraphs])
                    all_gecici.append(f"--- Belge (Geçici): {g_file.name} ---\n{text}")
                st.session_state.gecici_text = "\n\n".join(all_gecici)
            st.success("Geçici belgeler hafızaya alındı!")

    st.session_state.yonetmelik_text = permanent_text + "\n\n" + st.session_state.get("gecici_text", "")
    
    st.markdown("---")
    if st.button("Sonraki Adım: Tutanak Yükleme"):
        st.session_state.step = 2
        st.rerun()

# ----------------- ADIM 2: TUTANAK YÜKLEME -----------------
with st.expander("02  •  Tutanak Yükleme ve Analizi", expanded=(st.session_state.get('expand_all', False) or st.session_state.step == 2)):
    st.write("Lütfen olaya ilişkin tutanağı PDF veya Word formatında yükleyin.")
    tutanak_file = st.file_uploader("Tutanak Seç", type=["pdf", "docx"])
    
    if tutanak_file:
        try:
            if tutanak_file.type == "application/pdf":
                pdf_doc = fitz.open(stream=tutanak_file.getvalue(), filetype="pdf")
                extracted_text = "\n".join(page.get_text() for page in pdf_doc)
            else:
                extracted_text = "\n".join(p.text for p in Document(tutanak_file).paragraphs)
            st.session_state.tutanak_text = extracted_text
        except Exception as error:
            st.error(f"Tutanak metni okunamadı: {error}")
            extracted_text = ""

        st.caption("Tutanak metnini ve özet alanlarını gözden geçirerek kaydedin.")
        emp_name = st.text_input("Çalışan Adı Soyadı", value=st.session_state.get("employee_name", ""))
        emp_details = st.text_input("Sicil No / Unvan", value=st.session_state.get("employee_details", ""))
        summary = st.text_area("Tutanak Özeti", value=st.session_state.get("tutanak_summary", extracted_text), height=220)
        if st.button("Tutanağı Kaydet"):
            if not emp_name.strip() or not summary.strip():
                st.warning("Çalışan adı ve tutanak özeti zorunludur.")
            else:
                st.session_state.employee_name = emp_name.strip()
                st.session_state.employee_details = emp_details.strip()
                st.session_state.tutanak_summary = summary.strip()
                if not st.session_state.get('current_record_id'):
                    st.session_state.current_record_id = hr_db.create_record(emp_name.strip(), emp_details.strip(), "Tutanak Tutuldu")
                save_current_state_to_db()
                hr_db.log_action(st.session_state.username, f"Yeni tutanak kaydetti ve dosya açtı: {emp_name.strip()}")
                st.success("Tutanak kaydedildi.")
            
    if st.session_state.tutanak_summary:
        st.subheader("Tutanak Özeti")
        st.write(st.session_state.tutanak_summary)
        if st.button("Sonraki Adım: Savunma İstemi Oluştur"):
            st.session_state.step = 3
            st.rerun()

# ----------------- ADIM 3: SAVUNMA İSTEMİ -----------------
with st.expander("03  •  Savunma İstem Belgesi", expanded=(st.session_state.get('expand_all', False) or st.session_state.step == 3)):
    st.write("Tutanak özetine dayanarak çalışana verilecek resmi savunma istem metni oluşturuluyor.")
    
    if not st.session_state.tutanak_summary:
        st.warning("Lütfen önce 2. adımdan tutanak yükleyin.")
    else:
        st.caption("Savunma istem metnini kontrol ederek buraya girin.")
        st.session_state.savunma_istem_metni = st.text_area("Savunma İstem Metni", value=st.session_state.get("savunma_istem_metni", ""), height=300)
        if st.button("Savunma İstemini Kaydet"):
            if not st.session_state.savunma_istem_metni.strip():
                st.warning("Kaydetmek için metni girin.")
            else:
                if st.session_state.get('current_record_id'):
                    hr_db.update_status(st.session_state.current_record_id, "Savunma Bekleniyor")
                    save_current_state_to_db()
                hr_db.log_action(st.session_state.username, "Savunma istem belgesini kaydetti.")
                st.success("Savunma istem metni kaydedildi.")
                
        if st.session_state.savunma_istem_metni:
            st.text_area("Oluşturulan Metin", st.session_state.savunma_istem_metni, height=300)
            
            docx_file = create_docx(st.session_state.savunma_istem_metni, "Savunma İstem Yazısı")
            st.download_button("Word Olarak İndir (DOCX)", data=docx_file, file_name="savunma_istem_yazisi.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            if st.button("Sonraki Adım: Çalışan Savunmasını Yükle"):
                st.session_state.step = 4
                st.rerun()

# ----------------- ADIM 4: SAVUNMA YÜKLEME -----------------
with st.expander("04  •  Çalışan Savunması ve Değerlendirme", expanded=(st.session_state.get('expand_all', False) or st.session_state.step == 4)):
    st.write("Çalışanın el yazısı ile teslim ettiği PDF formatındaki savunmayı yükleyin.")
    savunma_file = st.file_uploader("Savunma Seç", type=["pdf"])
    
    if savunma_file:
        try:
            pdf_doc = fitz.open(stream=savunma_file.getvalue(), filetype="pdf")
            extracted_defense = "\n".join(page.get_text() for page in pdf_doc)
        except Exception as error:
            st.error(f"Savunma dosyası okunamadı: {error}")
            extracted_defense = ""
        st.caption("El yazısı OCR bu uygulamada çalıştırılmaz. Okunamayan kısımları çözümlenmiş metin olarak buraya girin.")
        st.session_state.savunma_text = st.text_area("Çalışan Savunması", value=st.session_state.get("savunma_text", extracted_defense), height=260)
        if st.button("Savunmayı Kaydet"):
            if not st.session_state.savunma_text.strip():
                st.warning("Kaydetmek için savunma metnini girin.")
            else:
                if st.session_state.get('current_record_id'):
                    hr_db.update_status(st.session_state.current_record_id, "Savunma Alındı")
                    save_current_state_to_db()
                hr_db.log_action(st.session_state.username, "Personele ait savunmayı kaydetti.")
                st.success("Savunma metni kaydedildi.")
            
    if st.session_state.savunma_text:
        st.subheader("Çalışan Savunması (Çözümlenmiş)")
        st.write(st.session_state.savunma_text)
        if st.button("Sonraki Adım: Hukuki Karar Motoru"):
            st.session_state.step = 5
            st.rerun()

# ----------------- ADIM 5: KARAR MOTORU -----------------
with st.expander("05  •  Nihai Karar Değerlendirmesi", expanded=(st.session_state.get('expand_all', False) or st.session_state.step == 5)):
    st.write("Şirket yönetmeliği, tutanak, savunma, İş Kanunu ve Yargıtay İçtihatları çerçevesinde değerlendirme yapılır.")
    
    st.caption("Değerlendirme notunuzu buraya girin. Nihai karar yetkili yönetici onayıyla verilir.")
    st.session_state.karar_sonucu = st.text_area("Karar Değerlendirme Notu", value=st.session_state.get("karar_sonucu", ""), height=300)
    if st.button("Değerlendirmeyi Kaydet"):
        if st.session_state.karar_sonucu.strip():
            save_current_state_to_db()
            hr_db.log_action(st.session_state.username, "Karar değerlendirme notunu kaydetti.")
            st.success("Değerlendirme notu kaydedildi.")
        else:
            st.warning("Kaydetmek için değerlendirme notunu girin.")
            
    if st.session_state.karar_sonucu:
        st.subheader("Karar Değerlendirme Notu ve Gerekçesi")
        st.write(st.session_state.karar_sonucu)
        
        st.markdown("### ✍️ Yönetici Kararı (Sizin Seçiminiz)")
        secenekler = ["İşlem Yapılmamasına (Savunmanın Kabulüne)", "Sözlü Uyarı", "Yazılı Uyarı", "Kınama", "Ücret Kesintisi", "İş Akdi Feshi (Tazminatlı)", "İş Akdi Feshi (Tazminatsız/Haklı Nedenle)"]
        secilen_karar = st.selectbox("Lütfen uygulanacak nihai aksiyonu seçin:", secenekler)
        
        if st.button("Kararı Onayla ve Belge Aşamasına Geç"):
            st.session_state.secilen_karar = secilen_karar
            if st.session_state.get('current_record_id'):
                hr_db.update_status(st.session_state.current_record_id, f"Sonuçlandı ({secilen_karar})")
                save_current_state_to_db()
            hr_db.log_action(st.session_state.username, f"Dosyayı karara bağladı: {secilen_karar}")
            st.session_state.step = 6
            st.rerun()

# ----------------- ADIM 6: KARAR ÇIKTISI -----------------
with st.expander("06  •  Resmî Karar Çıktısı", expanded=(st.session_state.get('expand_all', False) or st.session_state.step == 6)):
    st.write("Verilen karara uygun nihai bildirim metni.")
    
    if st.session_state.karar_sonucu:
        st.caption("Nihai bildirim metnini kontrol ederek buraya girin.")
        st.session_state.nihai_belge = st.text_area("Nihai Belge Metni", value=st.session_state.get("nihai_belge", ""), height=300)
        if st.button("Nihai Belgeyi Kaydet"):
            if st.session_state.nihai_belge.strip():
                save_current_state_to_db()
                hr_db.log_action(st.session_state.username, "Nihai belge metnini kaydetti.")
                st.success("Nihai belge kaydedildi.")
            else:
                st.warning("Kaydetmek için nihai belge metnini girin.")
                
    if st.session_state.nihai_belge:
        st.text_area("Nihai Belge", st.session_state.nihai_belge, height=300)
        default_baslik = f"{st.session_state.get('secilen_karar', 'Karar')} Bildirimi"
        belge_basligi = st.text_input("Belge Başlığı (Örn: Yazılı Uyarı, Kınama, Fesih Bildirimi)", value=default_baslik)
        docx_file = create_docx(st.session_state.nihai_belge, belge_basligi)
        st.download_button("Word Olarak İndir (DOCX)", data=docx_file, file_name="karar_bildirimi.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ----------------- ADIM 7: DOSYA TAKİBİ -----------------
with st.expander("07  •  Dosya ve Süreç Takibi", expanded=(st.session_state.get('expand_all', False) or st.session_state.step == 7)):
    st.write("Sistemdeki tüm disiplin dosyalarının güncel durumları.")
    
    records = hr_db.get_all_records()
    if not records:
        st.info("Henüz sisteme işlenmiş bir disiplin kaydı bulunmuyor.")
    else:
        for r in records:
            status = r['status']
            if status == "Tutanak Tutuldu":
                bg = "#F3F4F6"; text_color = "#4B5563"; border = "#E5E7EB"; icon = "📝"
            elif status == "Savunma Bekleniyor":
                bg = "#FFFBEB"; text_color = "#D97706"; border = "#FEF3C7"; icon = "⏳"
            elif status == "Savunma Alındı":
                bg = "#EFF6FF"; text_color = "#1D4ED8"; border = "#DBEAFE"; icon = "🛡️"
            else: # Sonuçlandı
                bg = "#ECFDF5"; text_color = "#065F46"; border = "#D1FAE5"; icon = "🟢"
            st.markdown("""<div style="background-color: #FFFFFF; border: 1px solid #E5E7EB; border-radius: 12px; padding: 15px 20px; margin-bottom: 12px; box-shadow: 0 1px 3px rgba(0,0,0,0.02);">""", unsafe_allow_html=True)
            col1, col2 = st.columns([5, 1])
            with col1:
                st.markdown(f"""
                <div style="display: flex; justify-content: space-between; align-items: center;">
                    <div>
                        <div style="font-size: 16px; font-weight: 700; color: #111827; margin-bottom: 4px;">{r['employee_name']}</div>
                        <div style="font-size: 13px; color: #6B7280;">{r['details']} • {r['file_number']}</div>
                    </div>
                    <div style="background-color: {bg}; color: {text_color}; padding: 6px 12px; border-radius: 20px; font-size: 13px; font-weight: 600; border: 1px solid {border}; display: flex; align-items: center;">
                        <span style="font-size: 12px; margin-right: 6px;">{icon}</span> {status}
                    </div>
                </div>
                """, unsafe_allow_html=True)
            with col2:
                if "Sonuçlandı" not in status:
                    has_state = 'state_json' in r.keys() and bool(r['state_json'])
                    btn_label = "Devam Et ➡️" if has_state else "Eski Kayıt ❌"
                    
                    if st.button(btn_label, key=f"devam_{r['id']}", disabled=not has_state, help="Sadece bu özellik eklendikten sonra oluşturulan dosyalar kaldığı yerden devam edebilir."):
                        import json
                        try:
                            saved_state = json.loads(r['state_json'])
                            st.session_state.current_record_id = r['id']
                            for k, v in saved_state.items():
                                st.session_state[k] = v
                            if status == "Tutanak Tutuldu": st.session_state.step = 3
                            elif status == "Savunma Bekleniyor": st.session_state.step = 4
                            elif status == "Savunma Alındı": st.session_state.step = 5
                            st.rerun()
                        except Exception as e:
                            st.error(f"Kayıt yüklenemedi: {e}")
                            
            st.markdown("</div>", unsafe_allow_html=True)
